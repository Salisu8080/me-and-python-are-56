"""
DOCX to PDF Converter - Robust script for converting Microsoft Word documents to PDFs

Requirements:
    pip install docx2pdf python-docx tqdm colorama PyPDF2 reportlab
    
Note: This script offers two conversion methods:
    1. Using docx2pdf (requires MS Word on Windows or LibreOffice on Linux/Mac)
    2. Using python-docx + reportlab (no external dependencies but limited formatting)

How to Use the Script

Simple usage - Convert a single file:
python docx_to_pdf.py document.docx

Specify output location:
python docx_to_pdf.py document.docx -o output.pdf

Convert all DOCX files in a directory:
python docx_to_pdf.py input_directory -o output_directory

Recursively process a directory:
python docx_to_pdf.py input_directory -r -o output_directory

Control parallel processing (e.g., use 4 workers):
python docx_to_pdf.py input_directory -w 4
"""

import os
import sys
import time
import argparse
import concurrent.futures
from pathlib import Path
import logging
from tqdm import tqdm
from colorama import init, Fore, Style

# Initialize colorama for cross-platform colored console output
init()

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("docx_to_pdf_conversion.log"),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# Try to import the required libraries
try:
    from docx2pdf import convert
    docx2pdf_available = True
except ImportError:
    docx2pdf_available = False
    logger.warning("docx2pdf library not found. Will use fallback conversion method.")

try:
    import docx
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    reportlab_available = True
except ImportError:
    reportlab_available = False
    if not docx2pdf_available:
        logger.error("Neither docx2pdf nor reportlab/python-docx are available. Please install required dependencies.")
        sys.exit(1)


def convert_with_docx2pdf(input_path, output_path=None):
    """Convert docx to pdf using docx2pdf library (requires MS Word or LibreOffice)"""
    try:
        if output_path:
            convert(input_path, output_path)
        else:
            output_path = str(input_path).replace('.docx', '.pdf')
            convert(input_path, output_path)
        return True, output_path
    except Exception as e:
        logger.error(f"Error converting {input_path} with docx2pdf: {str(e)}")
        return False, str(e)


def convert_with_reportlab(input_path, output_path=None):
    """Convert docx to pdf using python-docx and reportlab (fallback method)"""
    try:
        # If no output path is specified, create one
        if not output_path:
            output_path = str(input_path).replace('.docx', '.pdf')
        
        # Load the Word document
        doc = docx.Document(input_path)
        
        # Create a PDF document
        pdf = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Process document content
        styles = getSampleStyleSheet()
        flowables = []
        
        # Extract and add text from paragraphs
        for para in doc.paragraphs:
            if para.text:
                p = Paragraph(para.text, styles['Normal'])
                flowables.append(p)
                flowables.append(Spacer(1, 12))  # Add some space after each paragraph
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text for cell in row.cells)
                if row_text:
                    p = Paragraph(row_text, styles['Normal'])
                    flowables.append(p)
                    flowables.append(Spacer(1, 12))
        
        # Build the PDF
        pdf.build(flowables)
        
        return True, output_path
    except Exception as e:
        logger.error(f"Error converting {input_path} with reportlab: {str(e)}")
        return False, str(e)


def convert_docx_to_pdf(input_path, output_path=None, use_fallback=False):
    """Convert a single docx file to pdf using the best available method"""
    input_path = Path(input_path)
    
    # Validate input file
    if not input_path.exists():
        return False, f"Input file does not exist: {input_path}"
    
    if input_path.suffix.lower() != '.docx':
        return False, f"Input file is not a .docx file: {input_path}"
    
    # If output path is a directory, create full output path
    if output_path and os.path.isdir(output_path):
        output_path = os.path.join(output_path, input_path.stem + '.pdf')
    elif not output_path:
        output_path = str(input_path).replace('.docx', '.pdf')
    
    # Create output directory if it doesn't exist
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    
    # Choose conversion method
    if docx2pdf_available and not use_fallback:
        # Try docx2pdf first (better quality)
        success, result = convert_with_docx2pdf(str(input_path), output_path)
        if success:
            return True, result
        logger.warning(f"docx2pdf conversion failed, falling back to reportlab for {input_path}")
    
    if reportlab_available:
        # Use reportlab as fallback
        return convert_with_reportlab(str(input_path), output_path)
    
    return False, "No conversion method available"


def process_batch(files, output_dir=None, max_workers=None, use_fallback=False):
    """Process a batch of files with progress tracking"""
    if not max_workers:
        # Use half of available cores by default, but at least 1
        max_workers = max(1, (os.cpu_count() or 2) // 2)
    
    results = {"success": 0, "failed": 0, "failed_files": []}
    
    print(f"{Fore.CYAN}Starting conversion of {len(files)} files using {max_workers} workers...{Style.RESET_ALL}")
    
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        # Prepare all futures
        future_to_file = {
            executor.submit(convert_docx_to_pdf, file, output_dir, use_fallback): file
            for file in files
        }
        
        # Process as they complete with a progress bar
        for future in tqdm(concurrent.futures.as_completed(future_to_file), 
                          total=len(files), 
                          desc="Converting", 
                          unit="file"):
            file = future_to_file[future]
            try:
                success, result = future.result()
                if success:
                    results["success"] += 1
                    print(f"{Fore.GREEN}✓ Successfully converted: {file} -> {result}{Style.RESET_ALL}")
                else:
                    results["failed"] += 1
                    results["failed_files"].append((file, result))
                    print(f"{Fore.RED}✗ Failed to convert: {file} - {result}{Style.RESET_ALL}")
            except Exception as e:
                results["failed"] += 1
                results["failed_files"].append((file, str(e)))
                print(f"{Fore.RED}✗ Error processing: {file} - {str(e)}{Style.RESET_ALL}")
    
    return results


def find_docx_files(path):
    """Find all .docx files in the specified path (file or directory)"""
    path = Path(path)
    if path.is_file():
        if path.suffix.lower() == '.docx':
            return [str(path)]
        else:
            logger.warning(f"Skipping non-.docx file: {path}")
            return []
    elif path.is_dir():
        return [str(p) for p in path.glob('**/*.docx')]
    else:
        logger.error(f"Path does not exist: {path}")
        return []


def main():
    """Main function for command line usage"""
    parser = argparse.ArgumentParser(description='Convert DOCX files to PDF format')
    parser.add_argument('input', help='Input DOCX file or directory containing DOCX files')
    parser.add_argument('-o', '--output', help='Output PDF file or directory')
    parser.add_argument('-r', '--recursive', action='store_true', help='Recursively process directories')
    parser.add_argument('-f', '--fallback', action='store_true', help='Force use of fallback conversion method')
    parser.add_argument('-w', '--workers', type=int, help='Number of worker threads for batch processing')
    
    args = parser.parse_args()
    
    # Check conversion capabilities
    if args.fallback and not reportlab_available:
        logger.error("Fallback method requested but reportlab is not available")
        return 1
    
    if not docx2pdf_available and not reportlab_available:
        logger.error("No conversion methods available. Please install docx2pdf or reportlab.")
        return 1
    
    # Find all docx files
    input_path = Path(args.input)
    if input_path.is_file():
        files = [str(input_path)]
    else:
        if args.recursive:
            # Recursive search
            files = find_docx_files(input_path)
        else:
            # Non-recursive search
            files = [str(p) for p in input_path.glob('*.docx')]
    
    if not files:
        logger.warning(f"No DOCX files found in {args.input}")
        return 0
    
    # Process files
    start_time = time.time()
    
    if len(files) == 1 and input_path.is_file():
        # Single file conversion
        success, result = convert_docx_to_pdf(files[0], args.output, args.fallback)
        if success:
            print(f"{Fore.GREEN}Successfully converted: {files[0]} -> {result}{Style.RESET_ALL}")
            logger.info(f"Conversion successful: {files[0]} -> {result}")
        else:
            print(f"{Fore.RED}Failed to convert: {files[0]} - {result}{Style.RESET_ALL}")
            logger.error(f"Conversion failed: {files[0]} - {result}")
    else:
        # Batch processing
        results = process_batch(files, args.output, args.workers, args.fallback)
        
        # Report summary
        print(f"\n{Fore.CYAN}Conversion Complete!{Style.RESET_ALL}")
        print(f"Total files: {len(files)}")
        print(f"{Fore.GREEN}Successfully converted: {results['success']}{Style.RESET_ALL}")
        print(f"{Fore.RED}Failed conversions: {results['failed']}{Style.RESET_ALL}")
        
        if results["failed"] > 0:
            print(f"\n{Fore.YELLOW}Failed files:{Style.RESET_ALL}")
            for file, error in results["failed_files"]:
                print(f"  - {file}: {error}")
    
    # Final timing
    elapsed = time.time() - start_time
    print(f"\n{Fore.CYAN}Total time: {elapsed:.2f} seconds{Style.RESET_ALL}")
    
    return 0


if __name__ == "__main__":
    try:
        sys.exit(main())
    except KeyboardInterrupt:
        print(f"\n{Fore.YELLOW}Process interrupted by user.{Style.RESET_ALL}")
        sys.exit(1)
    except Exception as e:
        print(f"\n{Fore.RED}Unexpected error: {str(e)}{Style.RESET_ALL}")
        logger.exception("Unexpected error")
        sys.exit(1)