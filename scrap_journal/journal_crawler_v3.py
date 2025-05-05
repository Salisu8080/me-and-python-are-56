"""
pip install requests beautifulsoup4 tqdm pdfplumber python-docx re
python journal_crawler.py
"""

import requests
from bs4 import BeautifulSoup
import os
import json
import re
from pathlib import Path
from tqdm import tqdm
import concurrent.futures
from requests.adapters import HTTPAdapter
from requests.packages.urllib3.util.retry import Retry
import pdfplumber  # For PDF parsing
import docx  # For DOCX parsing


def sanitize_filename(filename):
    """Remove special characters that are invalid in filenames."""
    # Replace spaces with underscores and remove other invalid characters
    sanitized = re.sub(r'[\\/().,*?:"<>|]', "", filename)
    sanitized = re.sub(r'\s+', "_", sanitized)
    # Limit filename length
    if len(sanitized) > 100:
        sanitized = sanitized[:200]
    return sanitized


def create_directory(directory_path):
    """Create directory if it doesn't exist."""
    if not os.path.exists(directory_path):
        os.makedirs(directory_path)
        print(f"Created directory: {directory_path}")
    else:
        print(f"Directory already exists: {directory_path}")


def get_file_extension_from_url(url):
    """Extract file extension from URL."""
    # Check if URL has a file extension
    path = url.split('?')[0]  # Remove query parameters
    if '.' in path:
        ext = path.split('.')[-1].lower()
        # Return extension if it's pdf or docx
        if ext in ['pdf', 'docx']:
            return ext
    return None


def get_file_extension_from_content_type(content_type):
    """Determine file extension based on content-type header."""
    if content_type:
        content_type = content_type.lower()
        if 'application/pdf' in content_type:
            return 'pdf'
        elif 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' in content_type:
            return 'docx'
        elif 'application/msword' in content_type:
            return 'docx'
    return None


def download_file(url, save_path, session):
    """Download file from URL and save to specified path."""
    try:
        # Use the provided session with retries and headers already set
        with session.get(url, stream=True) as response:
            response.raise_for_status()  # Raise exception for HTTP errors
            
            # Get content type and determine file extension
            content_type = response.headers.get('content-type', '')
            url_ext = get_file_extension_from_url(url)
            content_ext = get_file_extension_from_content_type(content_type)
            
            # Determine file extension (prefer content-type, fallback to URL)
            file_ext = content_ext or url_ext or 'pdf'  # Default to pdf if can't determine
            
            # Update save path with correct extension if it doesn't match
            base_path, old_ext = os.path.splitext(save_path)
            if old_ext.lower() != f'.{file_ext}':
                save_path = f"{base_path}.{file_ext}"
            
            # Get file size for progress bar
            total_size = int(response.headers.get('content-length', 0))
            
            with open(save_path, 'wb') as file:
                # Download without tqdm progress bar for better performance in parallel downloads
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:  # filter out keep-alive new chunks
                        file.write(chunk)
            
            return {'success': True, 'path': save_path, 'url': url, 'extension': file_ext}
    except Exception as e:
        return {'success': False, 'path': save_path, 'url': url, 'error': str(e)}


def extract_abstract(input_value):
    """Clean and extract the abstract from the input value."""
    # Remove HTML tags
    soup = BeautifulSoup(input_value, 'html.parser')
    text = soup.get_text()
    # Normalize spaces
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def create_session():
    """Create a session with retry mechanism and browser-like headers."""
    session = requests.Session()
    
    # Set up retry strategy
    retry_strategy = Retry(
        total=3,
        backoff_factor=0.3,
        status_forcelist=[429, 500, 502, 503, 504],
        allowed_methods=["HEAD", "GET", "OPTIONS"]
    )
    
    adapter = HTTPAdapter(max_retries=retry_strategy, pool_connections=10, pool_maxsize=10)
    session.mount("http://", adapter)
    session.mount("https://", adapter)
    
    # Browser-like headers
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
        'Accept-Language': 'en-US,en;q=0.5',
        'Referer': 'https://www.soilsjournalnigeria.com/',
        'Connection': 'keep-alive',
        'Upgrade-Insecure-Requests': '1',
    }
    
    session.headers.update(headers)
    return session


def extract_text_from_pdf(pdf_path):
    """Extract all text from a PDF file."""
    try:
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n"
        return text
    except Exception as e:
        print(f"Error extracting text from PDF {pdf_path}: {str(e)}")
        return ""


def extract_text_from_docx(docx_path):
    """Extract all text from a DOCX file."""
    try:
        text = ""
        doc = docx.Document(docx_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Also get text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        
        return text
    except Exception as e:
        print(f"Error extracting text from DOCX {docx_path}: {str(e)}")
        return ""


def extract_emails_from_text(text):
    """Extract email addresses from text using regex."""
    # Pattern for email addresses
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    
    # Remove duplicates while preserving order
    unique_emails = []
    for email in emails:
        if email.lower() not in [e.lower() for e in unique_emails]:
            unique_emails.append(email)
    
    return unique_emails


def extract_keywords_from_text(text):
    """Extract keywords from academic paper text."""
    # Convert to lowercase for case-insensitive matching
    text_lower = text.lower()
    
    # Common patterns that precede keywords
    keyword_patterns = [
        r'keywords?\s*:(.+?)(?:\n\n|\n[A-Z]|\nabstract|\nintroduction)',
        r'key\s+words?\s*:(.+?)(?:\n\n|\n[A-Z]|\nabstract|\nintroduction)',
        r'indexing terms\s*:(.+?)(?:\n\n|\n[A-Z]|\nabstract|\nintroduction)'
    ]
    
    # Try each pattern
    for pattern in keyword_patterns:
        matches = re.search(pattern, text_lower, re.DOTALL | re.IGNORECASE)
        if matches:
            keywords_text = matches.group(1).strip()
            
            # Clean up the keywords
            # Keywords are typically separated by commas, semicolons, or line breaks
            keywords = re.split(r'[,;]\s*|\n', keywords_text)
            
            # Clean each keyword and filter out empty ones
            cleaned_keywords = [kw.strip() for kw in keywords if kw.strip()]
            
            return cleaned_keywords
    
    # If no matches found with the patterns, try to find any line containing "keywords:"
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if re.search(r'keywords?\s*:', line.lower()):
            # If found, the keywords are likely on this line or the next
            if ':' in line:
                keywords_text = line.split(':', 1)[1].strip()
                if keywords_text:
                    keywords = re.split(r'[,;]\s*', keywords_text)
                    return [kw.strip() for kw in keywords if kw.strip()]
            
            # Check the next line if this line doesn't contain the keywords
            if i + 1 < len(lines) and lines[i + 1].strip():
                keywords = re.split(r'[,;]\s*', lines[i + 1].strip())
                return [kw.strip() for kw in keywords if kw.strip()]
    
    return []


def process_file_for_metadata(file_path):
    """Extract metadata from PDF or DOCX file."""
    try:
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        # Extract text based on file type
        if ext == '.pdf':
            text = extract_text_from_pdf(file_path)
        elif ext == '.docx':
            text = extract_text_from_docx(file_path)
        else:
            print(f"Unsupported file type: {ext}")
            return {"email": "", "keywords": ""}
        
        # Extract emails and keywords
        emails = extract_emails_from_text(text)
        keywords = extract_keywords_from_text(text)
        
        return {
            "email": ", ".join(emails) if emails else "",
            "keywords": ", ".join(keywords) if keywords else ""
        }
    except Exception as e:
        print(f"Error processing metadata for {file_path}: {str(e)}")
        return {
            "email": "",
            "keywords": ""
        }


def main():
    # URL to process
    url = input("Please Enter the URL: ")
    folder = input("Please Enter folder name: ")
    max_workers = 8  # Number of concurrent downloads
    
    # Create save directory in Downloads folder
    downloads_dir = str(Path.home() / "Downloads")
    volume_dir = os.path.join(downloads_dir, folder)
    create_directory(volume_dir)
    
    # Create a session for connection pooling
    session = create_session()
    
    # Fetch and parse the webpage
    print(f"Fetching content from {url}...")
    try:
        response = session.get(url, timeout=30)
        response.raise_for_status()
    except requests.exceptions.HTTPError as e:
        print(f"HTTP Error: {e}")
        print("\nThe website is rejecting our request. Try accessing the site in your browser first.")
        return
    except requests.exceptions.ConnectionError:
        print("Connection Error: Failed to establish a connection to the website.")
        return
    except requests.exceptions.Timeout:
        print("Timeout Error: The request timed out. Try again later.")
        return
    except requests.exceptions.RequestException as e:
        print(f"Error fetching the webpage: {e}")
        return
    
    # Parse the page content
    soup = BeautifulSoup(response.text, 'html.parser')
    
    # Save the HTML for debugging
    debug_path = os.path.join(volume_dir, "page_source.html")
    with open(debug_path, 'w', encoding='utf-8') as f:
        f.write(response.text)
    print(f"Saved page source to {debug_path} for debugging")

    # Extract all elements by class in the order they appear in the document
    # This is a completely different approach that doesn't rely on DOM structure
    all_topics = []
    all_article_types = []
    all_pages = []
    all_authors = []
    all_downloads = []
    all_abstracts = []
    
    # Get all topics in document order
    for div in soup.find_all('div', class_='topic'):
        # Skip empty divs
        text = div.get_text(strip=True)
        if text:
            all_topics.append({
                'text': text,
                'element': div
            })
    
    # Get all article types in document order
    for div in soup.find_all('div', class_='article_type'):
        text = div.get_text(strip=True)
        all_article_types.append({
            'text': text,
            'element': div
        })
    
    # Get all pages in document order
    for div in soup.find_all('div', class_='pages'):
        text = div.get_text(strip=True)
        all_pages.append({
            'text': text,
            'element': div
        })
    
    # Get all authors in document order
    for div in soup.find_all('div', class_='authors'):
        text = div.get_text(strip=True)
        all_authors.append({
            'text': text,
            'element': div
        })
    
    # Get all downloads in document order
    for div in soup.find_all('div', class_='download'):
        a_tag = div.find('a', href=True)
        href = a_tag.get('href', '') if a_tag else ''
        all_downloads.append({
            'href': href,
            'element': div
        })
    
    # Get all abstracts in document order
    for input_tag in soup.find_all('input', type='hidden'):
        abstract_id = input_tag.get('id', '')
        abstract_value = input_tag.get('value', '')
        if abstract_id and abstract_value:
            all_abstracts.append({
                'id': abstract_id,
                'value': abstract_value,
                'element': input_tag
            })
    
    print(f"Found {len(all_topics)} topics, {len(all_article_types)} article types, {len(all_pages)} pages, " +
          f"{len(all_authors)} authors, {len(all_downloads)} downloads, and {len(all_abstracts)} abstracts.")
    
    # Check that we have a reasonable number of each
    if len(all_topics) < 1:
        print("Error: No topics found. Check the website structure.")
        return
    
    # Create metadata for each article based on the lists
    metadata = []
    download_tasks = []
    
    num_articles = len(all_topics)
    print(f"Processing {num_articles} articles...")
    
    # Print first few topics to verify they're different
    print("\nVerifying topics are unique:")
    for i, topic in enumerate(all_topics[:5]):
        print(f"Topic {i+1}: {topic['text'][:70]}...")
    
    # Process each article
    for i in range(num_articles):
        # Get article components, using index if available or empty default if not
        topic = all_topics[i]['text'] if i < len(all_topics) else f"Article {i+1}"
        article_type = all_article_types[i]['text'] if i < len(all_article_types) else "Original Research"
        pages = all_pages[i]['text'] if i < len(all_pages) else ""
        authors = all_authors[i]['text'] if i < len(all_authors) else ""
        download_link = all_downloads[i]['href'] if i < len(all_downloads) else ""
        
        # Get abstract if available
        abstract_id = ""
        abstract_text = ""
        if i < len(all_abstracts):
            abstract_id = all_abstracts[i]['id']
            abstract_text = extract_abstract(all_abstracts[i]['value'])
        
        print(f"\nProcessing article {i+1}/{num_articles}: {topic}")
        
        # Generate filename - use .pdf as a temporary extension, will be updated after download
        # based on the actual file type
        filename = f"{(i+1):02d}-{sanitize_filename(topic)}.pdf"
        file_path = os.path.join(volume_dir, filename)
        
        # Initialize article data with the new structure
        article_data = {
            "title": topic,                   # Renamed from "topic"
            "page_number": pages,             # Renamed from "pages"
            "authors": authors,
            "abstract": abstract_text,
            "file_path": file_path,          # Renamed from "filename"
            "keywords": "",                  # Will be populated after file processing
            "email": ""                      # Will be populated after file processing
        }
        
        # Add to metadata
        metadata.append(article_data)
        
        # If download link exists, prepare for download
        if download_link:
            # Make sure the download link is absolute
            if not download_link.startswith('http'):
                # Convert relative URL to absolute
                base_url = '/'.join(url.split('/')[:3])  # http://example.com
                if download_link.startswith('/'):
                    download_link = base_url + download_link
                else:
                    path_url = '/'.join(url.split('/')[:-1])  # http://example.com/path
                    download_link = path_url + '/' + download_link
            
            save_path = file_path
            download_tasks.append((article_data, save_path, download_link, i))
        else:
            print(f"Warning: No download link found for article: {topic}")
    
    # Use ThreadPoolExecutor for parallel downloads
    successful_downloads = []
    if download_tasks:
        print(f"\nDownloading {len(download_tasks)} files in parallel (max {max_workers} concurrent downloads)...")
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Start all download tasks
            future_to_article = {
                executor.submit(
                    download_file, 
                    download_link, 
                    save_path, 
                    session
                ): (article_data, i, idx) 
                for i, (article_data, save_path, download_link, idx) in enumerate(download_tasks)
            }
            
            # Process results as they complete
            for future in tqdm(concurrent.futures.as_completed(future_to_article), 
                              total=len(future_to_article),
                              desc="Downloading articles"):
                article_data, i, idx = future_to_article[future]
                try:
                    result = future.result()
                    if result['success']:
                        # Update the file path in metadata with the correct extension
                        actual_path = result['path']
                        metadata[idx]['file_path'] = actual_path
                        
                        # Print file extension for debugging
                        file_ext = os.path.splitext(actual_path)[1].lower()
                        print(f"Downloaded: {article_data['title']} ({file_ext})")
                        
                        successful_downloads.append((article_data, actual_path, idx))
                    else:
                        print(f"Error downloading {article_data['title']}: {result['error']}")
                except Exception as e:
                    print(f"Error processing {article_data['title']}: {str(e)}")
    else:
        print("No articles with download links found.")
    
    # Process downloaded files to extract email and keywords
    if successful_downloads:
        print("\nExtracting metadata from downloaded files...")
        for article_data, file_path, idx in tqdm(successful_downloads, desc="Processing files"):
            try:
                file_ext = os.path.splitext(file_path)[1].lower()
                print(f"Extracting metadata from: {os.path.basename(file_path)} ({file_ext})")
                
                # Process file based on its extension
                file_metadata = process_file_for_metadata(file_path)
                
                # Update metadata with information from the file
                metadata[idx]["email"] = file_metadata["email"]
                metadata[idx]["keywords"] = file_metadata["keywords"]
                
                print(f"  Keywords: {file_metadata['keywords'][:100]}...")
                print(f"  Email: {file_metadata['email']}")
            except Exception as e:
                print(f"Error processing file {file_path}: {str(e)}")
    
    # Save metadata to JSON file
    json_path = os.path.join(volume_dir, "metadata.json")
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(metadata, f, indent=2, ensure_ascii=False)
    
    print(f"\nComplete! Processed {len(metadata)} articles.")
    print(f"Downloaded {len(successful_downloads)} files ({sum(1 for _, path, _ in successful_downloads if path.lower().endswith('.pdf'))} PDFs, {sum(1 for _, path, _ in successful_downloads if path.lower().endswith('.docx'))} DOCXs)")
    print(f"Metadata saved to {json_path}")


if __name__ == "__main__":
    main()
